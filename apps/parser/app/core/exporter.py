"""
Word 导出器
将选中的题目导出为 Word 文档，支持水印
"""

from typing import List, Dict, Any, Optional
from pathlib import Path
from datetime import datetime
from collections import Counter
import copy
import io
import os
import re
import tempfile
import logging
import html as html_module
from urllib.parse import unquote

from PIL import Image
from bs4 import BeautifulSoup, NavigableString, Tag

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement, parse_xml

from .splitter import Question
from .wmf_converter import WMFConverter
from .latex_omml_export import latex_to_omml_element, should_try_omml, strip_latex_delimiters
from .convert_equations_cli import (
    get_convert_equations_exe,
    run_latex_to_mathtype_payload,
    decode_math_type_model,
)
from .mathtype_ole_embed import embed_mathtype_ole_in_paragraph, minimal_preview_png_bytes

logger = logging.getLogger(__name__)


def _export_diagnostics_to_doc_enabled() -> bool:
    """
    默认在 Word 里每题末尾写一行 [导出诊断]，便于排查串题/叠公式。
    正式卷面不需要时设环境变量：EXPORT_DIAGNOSTICS_TO_DOC=0（或 false/off/no）关闭。
    """
    raw = (os.environ.get("EXPORT_DIAGNOSTICS_TO_DOC") or "").strip().lower()
    if raw in ("0", "false", "no", "off"):
        return False
    return True


# PIL 不可用或读图失败时的回退宽度（英寸）
_FALLBACK_BLOCK_IN = 1.65
_FALLBACK_INLINE_IN = 0.64
_FALLBACK_OPTION_IN = 0.34
# 首选「宽×高」盒（英寸）；行内 max_h 不宜过大，否则瘦高图（△、竖分式）会压过汉字
_BOX_BLOCK = (1.75, 0.90)
# 行内公式图：略收紧盒高，减轻 Word 里「比正文大一圈」的观感
_BOX_INLINE = (1.28, 0.26)
_BOX_OPTION = (0.40, 0.205)
_BOX_DEFAULT = (1.36, 0.32)
# 为满足最小高度所需的宽度可超过盒宽，但不超过下列硬上限（避免撑破版心）
_HARD_MAX_W_BLOCK = 2.12
_HARD_MAX_W_INLINE = 2.72
_HARD_MAX_W_DEFAULT = 2.15
_HARD_MAX_W_OPTION = 0.43
# 各类别最小显示高度（英寸），与 12pt 汉字行高约对齐，避免公式像脚注
_MIN_DISPLAY_H_BLOCK = 0.138
_MIN_DISPLAY_H_INLINE = 0.124
_MIN_DISPLAY_H_OPTION = 0.064
_MIN_DISPLAY_H_DEFAULT = 0.115
# 最小显示宽度（英寸）
_MIN_W_BLOCK = 0.152
_MIN_W_INLINE = 0.142
_MIN_W_OPTION = 0.108
_MIN_W_DEFAULT = 0.132
# 行内图：像素「宽/高」较小时（瘦高符号），再压低显示高度上限，避免单符号过大
_INLINE_TALL_AR_CAP = 0.62
_INLINE_TALL_MAX_H_IN = 0.28

# 导出 docx 元数据（避免默认「python-docx」「Macintosh Word」等与资源管理器/属性页观感异常）
_DOC_META_AUTHOR = "题目导出"
_DOC_META_LAST_MODIFIED_BY = "题目导出"
_APP_XML_MAC_WORD = b"Microsoft Macintosh Word"
_APP_XML_WIN_WORD = b"Microsoft Office Word"


class WatermarkType:
    """水印类型"""
    TEXT = "text"
    IMAGE = "image"


class WordExporter:
    """Word 文档导出器"""
    
    def __init__(self):
        self.document = None
        self.image_dir: Optional[Path] = None
        # data/images，其下每个子目录为一个上传 file_id（与 HTML 里 /api/v1/images/{file_id}/ 一致）
        self.images_library_root: Optional[Path] = None
        self._embedded_basenames: set = set()
        # 已成功嵌入的图片绝对路径（补图去重用）；同路径允许多次嵌入 run，避免同文件多引用处空白
        self._embedded_resolved_paths: set = set()
        # 已成功处理过的 <img src>（规范化路径键），避免 DOM 嵌完后正则再往段尾插一遍
        self._embedded_img_src_keys: set = set()
        self._omml_fail_latex: set = set()
        # create_document 时记下首卷目录，便于每道题按 question.file_id 切换后仍能恢复默认
        self._export_image_dir_default: Optional[Path] = None
        self._export_images_library_root_default: Optional[Path] = None
        # 每题设置：file_id|question_id|题号，OMML 失败去重键前缀，减轻跨题同 LaTeX 误判
        self._export_formula_scope: str = ""

    def create_document(self, title: str = "", watermark_text: str = None, image_dir: Optional[str] = None) -> Document:
        """
        创建新文档
        
        Args:
            title: 文档标题
            watermark_text: 水印文字
            
        Returns:
            Document 对象
        """
        self.document = Document()
        self.image_dir = Path(image_dir) if image_dir else None
        self.images_library_root = self.image_dir.parent if self.image_dir else None
        self._export_image_dir_default = self.image_dir
        self._export_images_library_root_default = self.images_library_root
        self._embedded_basenames = set()
        self._embedded_resolved_paths = set()
        self._embedded_img_src_keys = set()
        self._omml_fail_latex = set()
        self._export_formula_scope = ""

        # 设置默认字体
        self._set_default_font()
        
        # 添加标题
        if title:
            self._add_title(title)
            
        # 添加水印
        if watermark_text:
            self._add_watermark(watermark_text)

        self._apply_document_metadata(title=title)
        return self.document

    def _apply_document_metadata(self, title: str = "") -> None:
        """覆盖 python-docx 默认属性，使 Windows 下作者/备注/程序名更接近 Word 本机保存。"""
        if not self.document:
            return
        cp = self.document.core_properties
        cp.author = _DOC_META_AUTHOR
        cp.last_modified_by = _DOC_META_LAST_MODIFIED_BY
        cp.title = (title or "").strip() or "导出的题目"
        cp.subject = ""
        cp.keywords = ""
        cp.category = ""
        cp.comments = ""
        now = datetime.now()
        cp.created = now
        cp.modified = now
        try:
            cp.revision = 1
        except Exception:
            pass
        self._patch_app_xml_application()

    def _refresh_document_modified_time(self) -> None:
        """保存前刷新修改时间。"""
        if self.document:
            self.document.core_properties.modified = datetime.now()

    def _patch_app_xml_application(self) -> None:
        """将 docProps/app.xml 中默认的 Macintosh Word 改为 Office Word（仅替换字节串）。"""
        if not self.document:
            return
        for part in self.document.part.package.iter_parts():
            if str(part.partname) != "/docProps/app.xml":
                continue
            blob = part.blob
            if not blob:
                break
            if _APP_XML_MAC_WORD in blob:
                part._blob = blob.replace(_APP_XML_MAC_WORD, _APP_XML_WIN_WORD)
            break

    def _set_default_font(self):
        """正文样式：中英混排 + 行距，与题干 12pt 运行一致，便于公式与汉字对齐。"""
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        try:
            pf = style.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = 1.2
        except Exception:
            pass
        
    def _add_title(self, title: str):
        """添加文档标题"""
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(title)
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        
        # 添加空行
        self.document.add_paragraph()
        
    def _add_watermark(self, text: str):
        """
        添加文字水印
        
        通过修改文档的 header 来实现水印效果
        """
        # 获取或创建 header
        section = self.document.sections[0]
        header = section.header
        
        # 创建水印的 XML
        watermark_xml = f'''
        <w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
             xmlns:v="urn:schemas-microsoft-com:vml"
             xmlns:o="urn:schemas-microsoft-com:office:office"
             xmlns:w10="urn:schemas-microsoft-com:office:word">
            <w:pict>
                <v:shapetype id="_x0000_t136" coordsize="21600,21600" o:spt="136" adj="10800" 
                    path="m@7,l@8,m@5,21600l@6,21600e">
                    <v:formulas>
                        <v:f eqn="sum #0 0 10800"/>
                        <v:f eqn="prod #0 2 1"/>
                        <v:f eqn="sum 21600 0 @1"/>
                        <v:f eqn="sum 0 0 @2"/>
                        <v:f eqn="sum 21600 0 @3"/>
                        <v:f eqn="if @0 @3 0"/>
                        <v:f eqn="if @0 21600 @1"/>
                        <v:f eqn="if @0 0 @2"/>
                        <v:f eqn="if @0 @4 21600"/>
                        <v:f eqn="mid @5 @6"/>
                        <v:f eqn="mid @8 @5"/>
                        <v:f eqn="mid @7 @8"/>
                        <v:f eqn="mid @6 @7"/>
                        <v:f eqn="sum @6 0 @5"/>
                    </v:formulas>
                    <v:path textpathok="t" o:connecttype="custom" o:connectlocs="@9,0;@10,10800;@11,21600;@12,10800" o:connectangles="270,180,90,0"/>
                    <v:textpath on="t" fitshape="t"/>
                    <v:handles>
                        <v:h position="#0,bottomRight" xrange="6629,14971"/>
                    </v:handles>
                    <o:lock v:ext="edit" text="t" shapetype="t"/>
                </v:shapetype>
                <v:shape id="PowerPlusWaterMarkObject" o:spid="_x0000_s2049" type="#_x0000_t136" 
                    style="position:absolute;margin-left:0;margin-top:0;width:500pt;height:100pt;rotation:315;z-index:-251658752;mso-position-horizontal:center;mso-position-horizontal-relative:margin;mso-position-vertical:center;mso-position-vertical-relative:margin"
                    o:allowincell="f" fillcolor="silver" stroked="f">
                    <v:fill opacity=".5"/>
                    <v:textpath style="font-family:&quot;宋体&quot;;font-size:60pt" string="{text}"/>
                    <w10:wrap anchorx="margin" anchory="margin"/>
                </v:shape>
            </w:pict>
        </w:r>
        '''
        
        # 添加水印到 header
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        run = paragraph.add_run()
        
        # 解析并插入水印 XML
        # python-docx 没有 OxmlElement.fromstring，这里使用 parse_xml。
        try:
            paragraph._p.append(parse_xml(watermark_xml))
        except Exception:
            # 水印失败不应阻断导出主流程
            pass

    @staticmethod
    def _snapshot_stem_export_segments(question: Question) -> List[Dict[str, Any]]:
        """本题 content_export_segments 深拷贝，避免多题导出时与快照/其它题共享引用导致串题。"""
        raw = getattr(question, "content_export_segments", None)
        if not isinstance(raw, list):
            return []
        out: List[Dict[str, Any]] = []
        for x in raw:
            if isinstance(x, dict):
                try:
                    out.append(copy.deepcopy(x))
                except Exception:
                    out.append(dict(x))
        return out

    def _segment_text_blob_has_all_option_labels(self, question: Question) -> bool:
        """
        判断 segments 文本里是否已包含本题全部选项标签（行首或换行后的 A. / B、形式）。
        仅当全部为真时才跳过 _add_options，避免片段未带选项时选择题「无选项」。
        """
        opts = getattr(question, "options", None) or []
        if not opts:
            return True
        segs = self._snapshot_stem_export_segments(question)
        parts: List[str] = []
        for s in segs:
            if self._export_segment_kind(s) == "text":
                parts.append(str(s.get("text") or ""))
        blob = "\n".join(parts)
        for opt in opts:
            lab = str(getattr(opt, "label", "") or "").strip()
            if not lab:
                return False
            pat = rf"(?:^|\n)\s*{re.escape(lab)}\s*[\.\、．]"
            if not re.search(pat, blob, re.MULTILINE):
                return False
        return True

    @staticmethod
    def _clone_oxml_fragment(el) -> Optional[Any]:
        """深拷贝 OMML 子树再挂到段落，避免 lxml 把同一节点从上一处移走导致跨段/跨题串公式。"""
        if el is None:
            return None
        try:
            from lxml import etree

            serialized = etree.tostring(el, encoding="unicode", with_tail=False)
            return parse_xml(serialized)
        except Exception as e:
            logger.debug("OMML 克隆失败，使用原节点: %s", e)
            return el

    def _write_question_stem(self, para, question: Question) -> bool:
        """
        只写本题题干：有 segments 时严格按快照顺序写 text / latex_inline / latex_block / image / paragraph_break；
        无 segments 时回退 content_html 或纯文本。不与其它题目混用数据。
        返回 True 表示已按 content_export_segments 写入；选项是否另写由 add_question 根据片段是否含选项标签决定。
        """
        segments = self._snapshot_stem_export_segments(question)
        logger.debug(
            "题干 id=%s number=%s 按段=%s 条数=%d",
            getattr(question, "id", "?"),
            getattr(question, "number", "?"),
            bool(segments),
            len(segments),
        )
        if segments:
            self._fill_stem_from_export_segments(
                para,
                segments,
                question.number,
                str(getattr(question, "id", "") or ""),
            )
            return True

        stem_html = (question.content_html or "").strip()
        if stem_html:
            stem_html = html_module.unescape(stem_html)
            stem_html = stem_html.replace('\\"', '"').replace("\\'", "'").replace("\\/", "/")
            root = BeautifulSoup(f'<div class="stem-root">{stem_html}</div>', "html.parser").select_one(
                ".stem-root"
            )
            if root:
                self._strip_leading_question_number(root, question.number)
                self._fill_stem_root(para, root)
            stem_has_img = bool(re.search(r"<img\b", stem_html, re.I))
            if not stem_has_img:
                self._add_missing_images_inline(para, question.images)
        else:
            self._add_formatted_text(para, question.content)
            self._add_question_images(question.images)
        if stem_html:
            logger.debug(
                "题干无 segments，已用 content_html：id=%s",
                getattr(question, "id", "?"),
            )
        return False

    def add_question(self, question: Question, include_answer: bool = False, 
                     include_analysis: bool = False):
        """
        添加一题到文档。
        与其它题隔离：每题入口清空嵌图去重集、ConvertEquations/OMML 缓存；按本题 file_id 切换 image_dir；
        题干仅用传入的 Question（建议调用方每题 deepcopy，避免多题共享嵌套 list/dict）。
        """
        # 每题重置「本题内已嵌图片」集合，避免跨题误伤；题干与选项共用本题集合
        self._embedded_basenames.clear()
        self._embedded_resolved_paths.clear()
        self._embedded_img_src_keys.clear()
        # 避免上一题的 ConvertEquations/OMML 状态影响本题；latex 字符串跨题偶发碰撞时尤甚
        self._omml_fail_latex.clear()
        # 跨卷组卷：每道题图片在 data/images/{该题 file_id}/；勿沿用首卷的 image_dir，否则 rglob 会嵌到别卷同名 wmf
        qfid = (getattr(question, "file_id", None) or "").strip()
        lib0 = self._export_images_library_root_default
        if qfid and lib0 is not None:
            self.image_dir = lib0 / qfid
            self.images_library_root = lib0
        else:
            self.image_dir = self._export_image_dir_default
            self.images_library_root = self._export_images_library_root_default
        qid = str(getattr(question, "id", "") or "").strip()
        qnum = getattr(question, "number", "")
        self._export_formula_scope = f"{qfid or '?'}|{qid or '?'}|n{qnum}"
        # 题号
        para = self.document.add_paragraph()
        run = para.add_run(f"{question.number}. ")
        run.font.bold = True
        run.font.size = Pt(12)

        # 题干：逐题只用本题快照，有 segments 则仅走按段写入（与 content_html 二选一）
        stem_from_segments = self._write_question_stem(para, question)

        # 选项：走 HTML 题干时必写；走 segments 时仅当片段文本能覆盖全部选项标签时才省略（否则补写，避免无选项）
        if question.options:
            if (not stem_from_segments) or (
                stem_from_segments
                and not self._segment_text_blob_has_all_option_labels(question)
            ):
                self._add_options(question.options)
            
        # 答案（与题干一致：$...$ 走 MathType OLE → OMML → 斜体回退）
        if include_answer and question.answer:
            para = self.document.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.3)
            run = para.add_run("【答案】")
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0, 128, 0)
            self._add_formatted_text(para, question.answer, font_size_pt=10.5)
            
        # 解析
        if include_analysis and question.analysis:
            para = self.document.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.3)
            run = para.add_run("【解析】")
            run.font.bold = True
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0, 0, 255)
            self._add_formatted_text(para, question.analysis, font_size_pt=10.5)

        if _export_diagnostics_to_doc_enabled():
            self._append_question_diagnostics_paragraph(question)

        # 添加空行
        self.document.add_paragraph()

    def _append_question_diagnostics_paragraph(self, question: Question) -> None:
        """每题一条：题干走 segments 还是 HTML、片段统计、目录与公式环境，便于分析串题/叠字根因。"""
        qfid = (getattr(question, "file_id", None) or "").strip()
        qid = str(getattr(question, "id", "") or "").strip()
        segs = self._snapshot_stem_export_segments(question)
        stem_html = (question.content_html or "").strip()
        hints: List[str] = []

        if segs:
            mode = "segments"
            kinds = Counter(self._export_segment_kind(s) for s in segs if isinstance(s, dict))
            kinds_s = ",".join(f"{k}:{kinds[k]}" for k in sorted(kinds.keys()) if k)
            text_dollar = 0
            api_in_fn = 0
            for s in segs:
                if not isinstance(s, dict):
                    continue
                if self._export_segment_kind(s) == "text" and "$" in str(s.get("text") or ""):
                    text_dollar += 1
                fn = str(s.get("filename") or s.get("source_image") or "")
                if "/api/v1/images/" in fn or fn.lower().startswith("http"):
                    api_in_fn += 1
            if text_dollar:
                hints.append(f"text段含$共{text_dollar}段(结构化下应无公式$)")
            if api_in_fn:
                hints.append(f"片段中URL式路径{api_in_fn}处(靠URL内file_id解析)")
            detail = f"seg={len(segs)} kinds[{kinds_s}] url字段={api_in_fn}"
        elif stem_html:
            mode = "content_html"
            detail = f"html_len={len(stem_html)}"
            hints.append("题干走content_html未走segments")
            if "$" in stem_html:
                hints.append("HTML含$符_add_formatted_text会拆公式")
        else:
            mode = "plain_content"
            detail = "无segments无html"
            hints.append("仅纯文本题干")

        if not qfid:
            hints.append("缺file_id→image_dir用导出首卷目录易同名串图")

        img_dir = self.image_dir
        dir_ok = bool(img_dir and img_dir.is_dir())
        if not dir_ok and qfid:
            hints.append(f"图片目录不存在:{img_dir}")

        ce = get_convert_equations_exe() is not None
        fpri = self._formula_priority()
        scope = getattr(self, "_export_formula_scope", "") or ""

        lib_root = self.images_library_root
        lib_s = str(lib_root) if lib_root else "(none)"

        hint_txt = " | ".join(hints) if hints else "(无)"
        line = (
            f"[导出诊断] id={qid or '?'} num={question.number} file_id={qfid or 'MISSING'} "
            f"stem_mode={mode} {detail} | img_dir_ok={int(dir_ok)} path={img_dir} | "
            f"lib_root={lib_s} | convert_equations={int(ce)} formula_pri={fpri} scope={scope} "
            f"|| 研判: {hint_txt}"
        )
        if len(line) > 1800:
            line = line[:1797] + "..."

        p = self.document.add_paragraph()
        r = p.add_run(line)
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(96, 96, 96)
        r.italic = True
        try:
            p.paragraph_format.left_indent = Inches(0.12)
            p.paragraph_format.space_before = Pt(2)
        except Exception:
            pass

    def _resolve_image_path(self, filename: str) -> Optional[Path]:
        """解析图片在本地磁盘中的真实路径"""
        if not self.image_dir:
            return None

        filename = (filename or "").strip().replace("\\", "/")
        if not filename:
            return None

        # 支持 filename 中可能带有 media/ 前缀
        candidates = [
            self.image_dir / filename,
            self.image_dir / "media" / filename,
            self.image_dir / Path(filename).name,
        ]
        if filename.startswith("media/"):
            candidates.append(self.image_dir / filename.replace("media/", "", 1))

        for p in candidates:
            if p.exists() and p.is_file():
                return p

        # 仅在本题 image_dir 下递归：避免整库 rglob 命中别卷同名文件（单机上多份 upload 时极常见）
        base = Path(filename).name
        if self.image_dir.is_dir() and base:
            try:
                for p in self.image_dir.rglob(base):
                    if p.is_file():
                        return p
            except OSError as e:
                logger.warning("rglob 图片失败: %s", e)

        # 旧行为：整库按文件名找（易串卷，默认关闭）。需要时再设 EXPORT_IMAGE_SEARCH_FULL_LIBRARY=1
        if (
            (os.environ.get("EXPORT_IMAGE_SEARCH_FULL_LIBRARY") or "").strip().lower()
            in ("1", "true", "yes")
            and self.images_library_root
            and self.images_library_root.is_dir()
            and base
        ):
            try:
                for p in self.images_library_root.rglob(base):
                    if p.is_file():
                        return p
            except OSError as e:
                logger.warning("库根 rglob 图片失败: %s", e)
        return None

    def _resolve_segment_asset_path(self, raw: str) -> Optional[Path]:
        """
        解析 segments 中的插图路径。
        若含 /api/v1/images/{file_id}/...（或绝对 URL），必须按 URL 内 file_id 定位，不能只用本题 image_dir + 文件名，
        否则跨卷/多题同名 media 文件会嵌到前一卷或前一题的图上（表现为公式串题）。
        """
        raw = (raw or "").strip()
        if not raw:
            return None
        if "/api/v1/images/" in raw or raw.lower().startswith("http"):
            pth = self._resolve_path_for_img_src(raw)
            if pth and pth.is_file():
                return pth
        p = self._resolve_image_path(raw)
        if p and p.is_file():
            return p
        base = Path(raw.replace("\\", "/")).name
        if base:
            p2 = self._resolve_image_path(base)
            if p2 and p2.is_file():
                return p2
        return None

    def _resolve_path_for_img_src(self, src: str) -> Optional[Path]:
        """按 /api/v1/images/{file_id}/filename 定位磁盘文件（与 Node 图片路由一致）"""
        src = self._normalize_img_src(src)
        m = re.search(r"(?:https?://[^/]+)?/api/v1/images/([^/]+)/([^?#]+)", src)
        if not m or not self.images_library_root:
            return None
        fid = m.group(1)
        fn = unquote(m.group(2).strip())
        sub = self.images_library_root / fid
        candidates = [
            sub / fn,
            sub / Path(fn).name,
            sub / "media" / Path(fn).name,
        ]
        for p in candidates:
            if p.is_file():
                return p
        name = Path(fn).name
        if sub.is_dir():
            try:
                for p in sub.rglob(name):
                    if p.is_file():
                        return p
            except OSError:
                pass
        return None

    def _add_question_images(self, images: List[str]):
        """将题目关联图片插入到文档"""
        if not images:
            return

        # 去重但保序，避免同图重复写入
        deduped = list(dict.fromkeys(images))
        for img in deduped:
            if Path(img).name in self._embedded_basenames:
                continue
            img_path = self._resolve_image_path(img)
            if not img_path:
                continue
            embed = self._embed_image_path_for_word(img_path)
            if not embed:
                continue
            try:
                p = self.document.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                w_in = self._picture_display_width_inches(embed, ['formula-image-block'])
                p.add_run().add_picture(str(embed), width=Inches(w_in))
                self._embedded_basenames.add(Path(img_path).name)
                try:
                    self._embedded_resolved_paths.add(str(img_path.resolve()))
                except OSError:
                    self._embedded_resolved_paths.add(str(img_path))
            except Exception:
                continue
            finally:
                self._cleanup_temp_embed(embed, img_path)
        
    def _img_src_dedupe_key(self, src: str) -> str:
        """同一图片 URL 不同写法（主机、转义）映射为同一键，供 src 级去重。"""
        s = self._normalize_img_src(src)
        if not s or s.startswith("data:"):
            return ""
        m = re.search(r"(?:https?://[^/]+)?(/api/v1/images/[^?#]+)", s)
        if m:
            try:
                return unquote(m.group(1).rstrip("/"))
            except Exception:
                return m.group(1).rstrip("/")
        return s.split("?")[0].strip()

    def _normalize_img_src(self, src: str) -> str:
        """消除 JSON/HTML 序列化产生的转义，便于解析真实 URL"""
        if not src:
            return ""
        s = src.strip()
        s = s.replace("\\/", "/")
        s = s.replace('\\"', '"').replace("\\'", "'")
        s = html_module.unescape(s)
        return s.strip().strip('"').strip("'")

    def _src_to_filename(self, src: str) -> Optional[str]:
        """从 /api/v1/images/{file_id}/xxx.png 等 URL 取出路径最后一段文件名"""
        src = self._normalize_img_src(src)
        if not src or src.startswith("data:"):
            return None
        # 支持绝对 URL: http://host/api/v1/images/...
        m = re.search(r"(?:https?://[^/]+)?/api/v1/images/[^/]+/([^?#]+)", src)
        if m:
            return unquote(m.group(1))
        try:
            return unquote(Path(src).name)
        except Exception:
            return None

    def _embed_image_path_for_word(self, path: Path) -> Optional[Path]:
        """python-docx 对 WMF 支持差，先转为 PNG 再嵌入"""
        if not path.exists():
            return None
        suf = path.suffix.lower()
        if suf not in ('.wmf', '.emf'):
            return path
        conv = WMFConverter()
        fd, tmp_png = tempfile.mkstemp(suffix='.png')
        os.close(fd)
        ok, _ = conv.convert(str(path), tmp_png)
        if ok and Path(tmp_png).exists():
            return Path(tmp_png)
        try:
            Path(tmp_png).unlink(missing_ok=True)
        except OSError:
            pass
        return None

    @staticmethod
    def _cleanup_temp_embed(embed: Path, original: Path) -> None:
        if embed != original and embed.exists():
            try:
                embed.unlink()
            except OSError:
                pass

    def _fallback_width_for_classes(self, classes: List[str]) -> float:
        cls = set(classes or [])
        if 'formula-image-block' in cls or 'math-block' in cls:
            return _FALLBACK_BLOCK_IN
        if 'option-image' in cls:
            return _FALLBACK_OPTION_IN
        if 'formula-image' in cls:
            return _FALLBACK_INLINE_IN
        return _FALLBACK_INLINE_IN

    @staticmethod
    def _box_floors_and_hard_cap(classes: List[str]):
        cls = set(classes or [])
        if 'formula-image-block' in cls or 'math-block' in cls:
            return _BOX_BLOCK, _MIN_DISPLAY_H_BLOCK, _MIN_W_BLOCK, _HARD_MAX_W_BLOCK
        if 'option-image' in cls:
            return _BOX_OPTION, _MIN_DISPLAY_H_OPTION, _MIN_W_OPTION, _HARD_MAX_W_OPTION
        if 'formula-image' in cls:
            return _BOX_INLINE, _MIN_DISPLAY_H_INLINE, _MIN_W_INLINE, _HARD_MAX_W_INLINE
        return _BOX_DEFAULT, _MIN_DISPLAY_H_DEFAULT, _MIN_W_DEFAULT, _HARD_MAX_W_DEFAULT

    def _picture_display_width_inches(self, embed_path: Path, classes: List[str]) -> float:
        """
        在盒内 contain；若高度低于 min_h，可放宽到 hard_max_w（解决长横式在盒宽下过扁）。
        瘦高位图（ar 小）额外压低允许高度，减轻 △、分式等单符号撑满一行的问题。
        """
        (max_w, max_h), min_h_floor, min_w_floor, hard_max_w = self._box_floors_and_hard_cap(
            classes
        )
        cls_set = set(classes or [])
        try:
            with Image.open(embed_path) as im:
                w_px, h_px = im.size
            if w_px < 1 or h_px < 1:
                return self._fallback_width_for_classes(classes)
            ar = w_px / float(h_px)
            # 块级整段公式不压；选项单独一套盒；行内/默认对瘦高图收紧高度上限
            mh = max_h
            if (
                'formula-image-block' not in cls_set
                and 'math-block' not in cls_set
                and 'option-image' not in cls_set
                and ar < _INLINE_TALL_AR_CAP
            ):
                mh = min(mh, _INLINE_TALL_MAX_H_IN)
            w_in = min(max_w, mh * ar)
            h_in = w_in / ar
            if h_in < min_h_floor:
                # 需要更宽才能达到 min_h；允许超过盒宽，直至 hard_max_w
                w_in = min(hard_max_w, min_h_floor * ar)
                h_in = w_in / ar
            if h_in > mh:
                w_in = mh * ar
                h_in = mh
            w_in = min(w_in, hard_max_w)
            if w_in < min_w_floor and (min_w_floor / ar) <= mh + 1e-9:
                w_in = min_w_floor
            return float(w_in)
        except Exception:
            return self._fallback_width_for_classes(classes)

    @staticmethod
    def _strip_leading_question_number(root: Tag, number: Any) -> None:
        """题干 HTML 若已含「1.」等与题号重复的前缀，去掉首处文本，避免出现「1. 1.」。"""
        if number is None or root is None:
            return
        try:
            n = int(number)
        except (TypeError, ValueError):
            return
        pat = re.compile(rf"^\s*{n}\s*[\.\、．]\s*")
        for node in root.descendants:
            if not isinstance(node, NavigableString):
                continue
            s = str(node)
            if not s.strip():
                continue
            if pat.match(s):
                node.replace_with(NavigableString(pat.sub("", s, count=1)))
            break

    def _add_picture_to_run(self, run, path: Path, classes: List[str]) -> None:
        try:
            key = str(path.resolve())
        except OSError:
            key = str(path)
        embed = self._embed_image_path_for_word(path)
        if not embed:
            return
        try:
            w = self._picture_display_width_inches(embed, classes)
            run.add_picture(str(embed), width=Inches(w))
            self._embedded_basenames.add(Path(path).name)
            self._embedded_resolved_paths.add(key)
        except Exception as e:
            logger.warning('嵌入图片失败: %s', e)
        finally:
            self._cleanup_temp_embed(embed, path)

    def _add_missing_images_inline(self, paragraph, images: List[str]) -> None:
        """按 question.images 补嵌尚未成功的文件（与 HTML/正则路径去重靠 _embedded_resolved_paths）。"""
        if not images or not (self.images_library_root or self.image_dir):
            return
        for img in dict.fromkeys(images):
            base = Path(img).name
            pth = self._resolve_image_path(img) or self._resolve_image_path(base)
            if not pth:
                continue
            try:
                k = str(pth.resolve())
            except OSError:
                k = str(pth)
            if k in self._embedded_resolved_paths:
                continue
            run = paragraph.add_run()
            self._add_picture_to_run(run, pth, ['formula-image'])

    def _extract_latex_from_math_element(self, element: Tag) -> str:
        raw = (element.get("data-latex") or "").strip()
        if raw:
            return strip_latex_delimiters(html_module.unescape(raw))
        inner = (element.get_text() or "").strip()
        return strip_latex_delimiters(inner)

    def _clear_paragraph_text_runs(self, paragraph) -> None:
        p_el = paragraph._p
        for child in list(p_el):
            if child.tag == qn("w:r"):
                p_el.remove(child)

    def _mark_image_path_as_embedded(self, filename_or_path: str) -> None:
        raw = (filename_or_path or "").strip()
        if not raw:
            return
        raw = html_module.unescape(raw.replace("\\/", "/").strip())
        pth: Optional[Path] = None
        if "/api/v1/images/" in raw or raw.startswith("http"):
            pth = self._resolve_path_for_img_src(raw)
        if not pth:
            pth = self._resolve_image_path(raw) or self._resolve_image_path(Path(raw).name)
        if not pth:
            return
        try:
            self._embedded_resolved_paths.add(str(pth.resolve()))
        except OSError:
            self._embedded_resolved_paths.add(str(pth))
        self._embedded_basenames.add(Path(pth).name)

    def _try_insert_math_omml(self, paragraph, element: Tag) -> bool:
        classes = element.get("class") or []
        if isinstance(classes, str):
            classes = [classes]
        is_block = any(c in ("math-block", "math-display") for c in classes)
        inline = not is_block
        latex = self._extract_latex_from_math_element(element)
        if not latex or not should_try_omml(latex):
            return False
        scope = getattr(self, "_export_formula_scope", "") or ""
        cache_key = f"{scope}|{'B' if is_block else 'I'}|{latex}"
        if cache_key in self._omml_fail_latex:
            return False
        el = latex_to_omml_element(latex, inline=inline)
        if el is None:
            self._omml_fail_latex.add(cache_key)
            return False
        el_use = self._clone_oxml_fragment(el)
        try:
            if inline:
                run = paragraph.add_run()
                run._element.append(el_use)
            else:
                self._clear_paragraph_text_runs(paragraph)
                paragraph._p.append(el_use)
            data_img = (element.get("data-image") or "").strip()
            if data_img:
                self._mark_image_path_as_embedded(data_img)
            return True
        except Exception as e:
            logger.debug("写入 OMML 到段落失败: %s", e)
            self._omml_fail_latex.add(cache_key)
            return False

    def _preview_png_for_mathtype_ole(self, wmf_b: Optional[bytes]) -> bytes:
        if not wmf_b or len(wmf_b) <= 8:
            return minimal_preview_png_bytes()
        tmp_wmf = None
        tmp_png = None
        try:
            conv = WMFConverter()
            tw = tempfile.NamedTemporaryFile(suffix=".wmf", delete=False)
            tw.write(wmf_b)
            tw.flush()
            tw.close()
            tmp_wmf = tw.name
            tp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
            tp.close()
            tmp_png = tp.name
            ok, _ = conv.convert(tmp_wmf, tmp_png)
            if ok and Path(tmp_png).is_file():
                png = Path(tmp_png).read_bytes()
                if png:
                    return png
        except Exception as e:
            logger.debug("WMF→PNG 预览失败，使用占位 PNG: %s", e)
        finally:
            for p in (tmp_wmf, tmp_png):
                if p:
                    try:
                        Path(p).unlink(missing_ok=True)
                    except OSError:
                        pass
        return minimal_preview_png_bytes()

    def _emit_latex_plain_fallback(
        self, paragraph, latex: str, *, is_block: bool, font_size_pt: float = 12
    ) -> None:
        """OMML/OLE/插图均失败时的最后回退：不再调用 ConvertEquations，避免重复插入与 $$ 拆分错乱。"""
        t = (latex or "").strip()
        if not t:
            return
        run = paragraph.add_run(f" {t} ")
        run.italic = True
        run.font.size = Pt(font_size_pt)
        if is_block:
            try:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                pass

    @staticmethod
    def _formula_priority() -> str:
        """
        导出公式（运行 Parser 的进程读取 EXPORT_FORMULA_PRIORITY）：
        - mathtype（默认）：LaTeX→MathType（OLE，失败则同次输出的 WMF 图）→ 仍失败再 Word OMML（可编辑降级）
        - omml：OMML 优先，再 MathType
        - auto：能发现 ConvertEquations.exe 则同 mathtype，否则同 omml
        """
        raw = (os.environ.get("EXPORT_FORMULA_PRIORITY") or "mathtype").strip().lower()
        if raw in ("word", "native", "omml"):
            return "omml"
        if raw == "auto":
            return "mathtype" if get_convert_equations_exe() is not None else "omml"
        return "mathtype"

    def _try_formula_engines_order(self, paragraph, element: Tag) -> bool:
        """默认先 MathType 再 OMML；仅当 EXPORT_FORMULA_PRIORITY=omml 时相反。"""
        pri = self._formula_priority()
        if pri == "omml":
            if self._try_insert_math_omml(paragraph, element):
                return True
            return self._try_insert_mathtype_via_convert_equations(paragraph, element)
        if self._try_insert_mathtype_via_convert_equations(paragraph, element):
            return True
        return self._try_insert_math_omml(paragraph, element)

    def _try_insert_mathtype_via_convert_equations(self, paragraph, element: Tag) -> bool:
        if get_convert_equations_exe() is None:
            return False
        latex = self._extract_latex_from_math_element(element)
        if not latex:
            return False
        classes = element.get("class") or []
        if isinstance(classes, str):
            classes = [classes]
        is_block = any(c in ("math-block", "math-display") for c in classes)
        data_img = (element.get("data-image") or "").strip()
        # 不在此缓存 ConvertEquations 结果：同 LaTeX 在不同题/不同位置须各自嵌入独立 OLE/关系，
        # 复用 payload 曾导致 docx 内预览图部件去重后多处以同图指向不同语义，表现为公式串位。
        data, err = run_latex_to_mathtype_payload(latex)
        if not data:
            logger.debug("ConvertEquations 未返回数据: %s", err)
            return False
        ole_b, wmf_b = decode_math_type_model(data)
        if ole_b:
            ole_b = bytes(ole_b)
        if wmf_b:
            wmf_b = bytes(wmf_b)
        img_cls = ["formula-image-block"] if is_block else ["formula-image"]

        # 始终优先 MathType OLE（可编辑），失败再 WMF 图；与 EXPORT_FORMULA_PRIORITY=mathtype 一致
        if ole_b:
            prog_id = (os.environ.get("EXPORT_MATHTYPE_OLE_PROGID") or "").strip() or None
            png_b = self._preview_png_for_mathtype_ole(wmf_b)
            ok = embed_mathtype_ole_in_paragraph(
                paragraph,
                ole_b,
                png_b,
                prog_id=prog_id,
                is_block=is_block,
            )
            if ok:
                if data_img:
                    self._mark_image_path_as_embedded(data_img)
                logger.debug("已内嵌 MathType OLE（%d 字节）", len(ole_b))
                return True
            logger.debug("MathType OLE 嵌入失败，尝试 WMF 回退")

        if wmf_b and len(wmf_b) > 8:
            tmp_wmf: Optional[str] = None
            try:
                tw = tempfile.NamedTemporaryFile(suffix=".wmf", delete=False)
                tw.write(wmf_b)
                tw.close()
                tmp_wmf = tw.name
                run = paragraph.add_run()
                self._add_picture_to_run(run, Path(tmp_wmf), img_cls)
                if data_img:
                    self._mark_image_path_as_embedded(data_img)
                logger.debug("ConvertEquations 无 OLE 或 OLE 失败，已嵌入 WMF 转图（%d 字节）", len(wmf_b))
                return True
            except Exception as e:
                logger.debug("WMF 嵌入失败: %s", e)
                return False
            finally:
                if tmp_wmf:
                    try:
                        Path(tmp_wmf).unlink(missing_ok=True)
                    except OSError:
                        pass
        return False

    def _fill_stem_root(self, para, root: Tag) -> None:
        """题干根节点：顶层多个 <p> / 块级 <div> 分段落，题号与首段正文同一段。"""
        first_top = True
        # 当前承接「流式正文」的段落；块级公式另起段后，后续顶层文本不得写回题号段，否则与公式顺序错乱
        current_para = para
        start_new_para_before_next_flow = False

        for child in root.children:
            if isinstance(child, NavigableString):
                s = str(child)
                if s.strip():
                    if start_new_para_before_next_flow:
                        current_para = self.document.add_paragraph()
                        start_new_para_before_next_flow = False
                    self._add_formatted_text(current_para, s)
                    first_top = False
                continue
            if not isinstance(child, Tag):
                continue
            if child.name == 'p':
                target = para if first_top else self.document.add_paragraph()
                self._fill_inline(target, child)
                current_para = target
                first_top = False
                start_new_para_before_next_flow = True
            elif child.name == 'div':
                cl = child.get('class') or []
                if isinstance(cl, str):
                    cl = [cl]
                is_center = any(
                    x in ('math-block', 'math-display', 'formula-image-block') for x in cl
                )
                # 块级居中公式单独成段，避免把题号所在段整体居中
                if is_center:
                    target = self.document.add_paragraph()
                    target.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    self._fill_inline(target, child)
                    start_new_para_before_next_flow = True
                elif first_top:
                    target = para
                    self._fill_inline(target, child)
                    current_para = target
                    start_new_para_before_next_flow = False
                else:
                    target = self.document.add_paragraph()
                    self._fill_inline(target, child)
                    current_para = target
                    start_new_para_before_next_flow = False
                first_top = False
            else:
                if start_new_para_before_next_flow:
                    current_para = self.document.add_paragraph()
                    start_new_para_before_next_flow = False
                self._fill_inline(current_para, child)
                first_top = False
        try:
            para.paragraph_format.space_after = Pt(4)
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        except Exception:
            pass

    def _strip_leading_number_first_text_segment(
        self, segments: List[Dict[str, Any]], number: Any
    ) -> None:
        """
        去掉与 exporter 题号 run 重复的「n.」前缀。
        题号常被拆成多段 text（如 \"2 \" + \". 已知 \"），须合并后再匹配；数字与点之间允许空格。
        """
        if number is None:
            return
        try:
            n = int(number)
        except (TypeError, ValueError):
            return
        pat = re.compile(rf"^\s*{n}\s*[\.\、．]\s*")
        idxs: List[int] = []
        for i, seg in enumerate(segments):
            if self._export_segment_kind(seg) != "text":
                break
            idxs.append(i)
        if not idxs:
            return
        combined = "".join(str(segments[j].get("text", "")) for j in idxs)
        if not pat.match(combined):
            return
        rest = pat.sub("", combined, count=1)
        segments[idxs[0]]["text"] = rest
        for j in idxs[1:]:
            segments[j]["text"] = ""

    def _run_has_picture_or_object(self, run) -> bool:
        try:
            ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            el = run._element
            return bool(
                el.findall(f".//{{{ns}}}drawing") or el.findall(f".//{{{ns}}}object")
            )
        except Exception:
            return False

    def _paragraph_has_visual_content(self, paragraph) -> bool:
        if not paragraph:
            return False
        if (paragraph.text or "").strip():
            return True
        for run in paragraph.runs:
            if self._run_has_picture_or_object(run):
                return True
        return False

    def _synthetic_math_block_div(self, latex: str) -> Optional[Tag]:
        """与 content_html 块级公式同结构，供 _try_formula_engines_order 使用。"""
        raw = (latex or "").strip()
        if not raw:
            return None
        attr = html_module.escape(raw, quote=True)
        body = self._escape_html_body(raw)
        frag = f'<div class="math-block" data-latex="{attr}">$${body}$$</div>'
        soup = BeautifulSoup(frag, "html.parser")
        return soup.select_one("div.math-block")

    def _emit_segment_latex_inline(self, paragraph, seg: Dict[str, Any]) -> None:
        latex = str(seg.get("latex") or "").strip()
        src_img = str(seg.get("source_image") or "").strip()
        if not latex and not src_img:
            return
        # 与 _fill_inline 一致：优先 LaTeX→MathType/OMML（可编辑），失败再用原卷 source_image，最后斜体 LaTeX。
        # data-image 参与 ConvertEquations 缓存键，减轻错误复用 OLE。
        span: Optional[Tag] = None
        if latex:
            latex_escaped = html_module.escape(latex, quote=True)
            body = self._escape_html_body(latex)
            img_attr = ""
            if src_img:
                img_attr = f' data-image="{html_module.escape(src_img, quote=True)}"'
            frag = f'<span class="math-inline" data-latex="{latex_escaped}"{img_attr}>${body}$</span>'
            span = BeautifulSoup(frag, "html.parser").select_one("span.math-inline")
        if span is not None and self._try_formula_engines_order(paragraph, span):
            return
        if src_img:
            pth = self._resolve_segment_asset_path(src_img)
            if pth and pth.is_file():
                run = paragraph.add_run()
                self._add_picture_to_run(run, pth, ["formula-image"])
                return
        if span is not None:
            latex_fb = self._extract_latex_from_math_element(span)
            if latex_fb:
                self._emit_latex_plain_fallback(
                    paragraph, latex_fb, is_block=False, font_size_pt=12
                )

    @staticmethod
    def _export_segment_kind(seg: Dict[str, Any]) -> str:
        """兼容 kind / type 及常见驼峰写法，避免静默跳过整段导致顺序错乱。"""
        raw = seg.get("kind")
        if raw is None or (isinstance(raw, str) and not raw.strip()):
            raw = seg.get("type")
        if not isinstance(raw, str):
            return ""
        s = raw.strip().lower().replace("-", "_")
        aliases = {
            "latexinline": "latex_inline",
            "latexblock": "latex_block",
            "paragraphbreak": "paragraph_break",
        }
        return aliases.get(s, s)

    @classmethod
    def _normalize_export_segment_dict(cls, seg: Dict[str, Any]) -> Dict[str, Any]:
        """
        统一片段字段：JSON/前端可能用 type、content、驼峰等，与 splitter 的 kind/text 对齐。
        """
        s = dict(seg)
        if (s.get("kind") is None or (isinstance(s.get("kind"), str) and not str(s.get("kind")).strip())) and s.get(
            "type"
        ) is not None:
            s["kind"] = s.get("type")
        # 正文：content_items 用 content，export 用 text
        if s.get("text") is None and s.get("content") is not None:
            s["text"] = s.get("content")
        if s.get("latex") is None:
            for key in ("latexContent", "latex_content", "tex"):
                if s.get(key) is not None:
                    s["latex"] = s[key]
                    break
        if not s.get("source_image"):
            alt = s.get("sourceImage") or s.get("source_image_path")
            if alt:
                s["source_image"] = alt
        if not s.get("filename"):
            for key in ("path", "src", "file", "imageFilename", "image_filename"):
                v = s.get(key)
                if v:
                    s["filename"] = str(v).strip()
                    break
        return s

    @staticmethod
    def _export_debug_segments_enabled() -> bool:
        """设置 EXPORT_DEBUG_SEGMENTS=1 时，在日志中按索引打印每条 segment，便于核对是否走按段导出。"""
        return (os.environ.get("EXPORT_DEBUG_SEGMENTS") or "").strip().lower() in (
            "1",
            "true",
            "yes",
        )

    @staticmethod
    def _export_segment_formula_dedupe_enabled() -> bool:
        """
        默认关闭相邻公式片段去重：仅用 basename 等指纹易把两道不同图/式判成重复而跳过，题干出现漏空或「窜位」观感。
        若解析确会输出紧邻重复 latex_inline 导致 Word 叠字，再设 EXPORT_SEGMENT_DEDUPE_FORMULAS=1。
        """
        return (os.environ.get("EXPORT_SEGMENT_DEDUPE_FORMULAS") or "").strip().lower() in (
            "1",
            "true",
            "yes",
        )

    def _add_plain_segment_text(self, paragraph, text: str, font_size_pt: float = 12) -> None:
        """
        按段导出时的正文：直接写入，不按 $ 拆分。
        segments 与 content_items 同源，公式已由 latex_inline / latex_block / image 表达；
        若此处再 _add_formatted_text，会把正文里残留的 $...$ 与结构化公式各插一遍，多题时表现为同一式子叠多层。
        """
        t = self._strip_word_equation_text_noise(text or "")
        if not t:
            return
        run = paragraph.add_run(t)
        run.font.size = Pt(font_size_pt)

    @staticmethod
    def _strip_word_equation_text_noise(s: str) -> str:
        """
        Word 解析后纯文本里常残留公式占位符 U+25A1（□），如「□P」「□(线性式)」。
        不做线性式解析，仅去掉明显占位，减轻与 latex_inline 叠显的怪异感。
        """
        if not s:
            return s
        box = "\u25a1"
        out = s
        out = re.sub(re.escape(box) + r"\s*\(", "(", out)
        out = re.sub(re.escape(box) + r"(?=[A-Za-z])", "", out)
        out = out.replace(box, "")
        return out

    @staticmethod
    def _norm_latex_for_dedupe(latex: str) -> str:
        """同一式子若仅空白/换行不同，不应重复嵌多遍 OLE。"""
        t = (latex or "").strip()
        t = re.sub(r"\s+", " ", t)
        return t

    @classmethod
    def _segment_formula_fingerprint(cls, seg: Dict[str, Any], kind: str) -> tuple:
        if kind == "latex_inline":
            return (
                "latex_inline",
                cls._norm_latex_for_dedupe(str(seg.get("latex") or "")),
                str(seg.get("source_image") or "").strip(),
            )
        if kind == "latex_block":
            return ("latex_block", cls._norm_latex_for_dedupe(str(seg.get("latex") or "")))
        if kind == "image":
            fn = str(seg.get("filename") or "").strip().replace("\\", "/")
            # 勿只用 basename：同卷内多图可能同名引用或不同子路径，误判重复会吞掉整段图/OLE
            return ("image", fn, str(seg.get("display") or "inline"))
        return (kind,)

    @classmethod
    def _dedupe_repeated_formula_segments(
        cls, segs: List[Dict[str, Any]]
    ) -> List[Dict[str, Any]]:
        """
        去掉重复的公式类片段：紧邻重复，或中间仅空白 text 隔开（解析常把同一式子拆成多条 latex_inline，
        Word 里会叠多层）。遇非空正文或 paragraph_break 则重置。
        """
        out: List[Dict[str, Any]] = []
        last_fp: Optional[tuple] = None
        for seg in segs:
            k = cls._export_segment_kind(seg)
            if k == "text":
                if not str(seg.get("text") or "").strip():
                    continue
                last_fp = None
                out.append(seg)
                continue
            if k == "paragraph_break":
                last_fp = None
                out.append(seg)
                continue
            if k in ("latex_inline", "latex_block", "image"):
                fp = cls._segment_formula_fingerprint(seg, k)
                if fp == last_fp:
                    continue
                last_fp = fp
                out.append(seg)
                continue
            last_fp = None
            out.append(seg)
        return out

    def _fill_stem_from_export_segments(
        self,
        para_with_number,
        segments: List[Dict[str, Any]],
        question_number: Any,
        question_id: str = "",
    ) -> None:
        """按解析阶段 content_export_segments 顺序写 Word，避免再依赖拼接后的 content_html DOM。"""
        if not segments:
            return
        segs = [self._normalize_export_segment_dict(dict(s)) for s in segments]
        self._strip_leading_number_first_text_segment(segs, question_number)
        if self._export_segment_formula_dedupe_enabled():
            segs = self._dedupe_repeated_formula_segments(segs)
        if self._export_debug_segments_enabled():
            logger.info(
                "[EXPORT_DEBUG_SEGMENTS] 题干按段导出 question_id=%s number=%s 条数=%d",
                question_id or "?",
                question_number,
                len(segs),
            )
        current = para_with_number
        for i, seg in enumerate(segs):
            k = self._export_segment_kind(seg)
            if self._export_debug_segments_enabled():
                if k == "text":
                    pv = repr((seg.get("text") or "")[:100])
                elif k in ("latex_inline", "latex_block"):
                    pv = repr((seg.get("latex") or "")[:100])
                elif k == "image":
                    pv = repr((seg.get("filename") or "")[:100])
                elif k == "paragraph_break":
                    pv = "paragraph_break"
                else:
                    pv = repr(seg)[:120]
                logger.info("  [%d] kind=%s %s", i, k, pv)
            if k == "paragraph_break":
                if self._paragraph_has_visual_content(current):
                    current = self.document.add_paragraph()
                continue
            if k == "text":
                t = seg.get("text") or ""
                if t:
                    self._add_plain_segment_text(current, t, 12)
                continue
            if k == "latex_inline":
                self._emit_segment_latex_inline(current, seg)
                continue
            if k == "latex_block":
                blk = self.document.add_paragraph()
                blk.alignment = WD_ALIGN_PARAGRAPH.CENTER
                div = self._synthetic_math_block_div(str(seg.get("latex") or ""))
                if div:
                    self._try_formula_engines_order(blk, div)
                current = self.document.add_paragraph()
                continue
            if k == "image":
                fn = str(seg.get("filename") or "").strip()
                if not fn:
                    continue
                pth = self._resolve_segment_asset_path(fn)
                if not pth:
                    logger.warning(
                        "导出 segments 找不到图片: fn=%r image_dir=%s lib=%s",
                        fn,
                        self.image_dir,
                        self.images_library_root,
                    )
                    continue
                disp = seg.get("display") or "inline"
                if disp == "block":
                    blk = self.document.add_paragraph()
                    blk.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = blk.add_run()
                    self._add_picture_to_run(run, pth, ["formula-image-block"])
                    current = self.document.add_paragraph()
                else:
                    run = current.add_run()
                    self._add_picture_to_run(run, pth, ["formula-image"])
                continue
            logger.debug("跳过未知 content_export_segments 项: kind=%r seg=%r", k, seg)
        try:
            para_with_number.paragraph_format.space_after = Pt(4)
            para_with_number.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        except Exception:
            pass

    def _emit_img_tag(self, paragraph, element: Tag) -> None:
        """嵌入单个 <img>（含作为遍历根节点的 img：无子节点，必须在 _fill_inline 入口单独处理）。"""
        src = element.get('src') or ''
        cls = element.get('class') or []
        if isinstance(cls, str):
            cls = [cls]
        fn = self._src_to_filename(src)
        if fn or src:
            pth = self._resolve_path_for_img_src(src) or (
                self._resolve_image_path(fn) if fn else None
            )
            if pth:
                run = paragraph.add_run()
                self._add_picture_to_run(run, pth, cls)
                sk = self._img_src_dedupe_key(src)
                if sk:
                    self._embedded_img_src_keys.add(sk)
            else:
                logger.warning(
                    "导出找不到图片文件: src=%r -> fn=%r image_dir=%s lib=%s",
                    src[:200] if src else "",
                    fn,
                    self.image_dir,
                    self.images_library_root,
                )
        else:
            logger.warning("导出无法从 img 解析 src: %r", (element.get("src") or "")[:200])

    def _fill_inline(self, paragraph, element: Tag) -> None:
        """将 HTML 片段写入段落：文本 + 行内公式图（按子节点顺序，与 content_html 一致）。"""
        # <img> 无子节点；若仅作为递归根传入，上面 for 不会执行，导致图片丢失、顺序后移
        if element.name == "img":
            self._emit_img_tag(paragraph, element)
            return
        # math-inline/math-block：MathType（默认）/ OMML（可切换）→ data-image → 斜体回退
        if element.name in ("span", "div"):
            classes = element.get("class") or []
            if isinstance(classes, str):
                classes = [classes]
            is_math_node = any(c in ("math-inline", "math-block", "math-display") for c in classes)
            if is_math_node:
                is_block_math = any(c in ("math-block", "math-display") for c in classes)
                if self._try_formula_engines_order(paragraph, element):
                    return
                data_image = (element.get("data-image") or "").strip()
                if data_image:
                    pth = self._resolve_image_path(data_image)
                    if pth:
                        run = paragraph.add_run()
                        img_classes = (
                            ["formula-image-block"]
                            if "math-block" in classes or "math-display" in classes
                            else ["formula-image"]
                        )
                        self._add_picture_to_run(run, pth, img_classes)
                        return
                # 禁止再遍历子节点：子节点中的 $$...$$ 经 _add_formatted_text 的 split('$') 会破坏版式，
                # 且可能与已失败的 OMML 组合成「同一式子多次占位、巨型叠字」。
                latex_fb = self._extract_latex_from_math_element(element)
                if latex_fb:
                    self._emit_latex_plain_fallback(
                        paragraph, latex_fb, is_block=is_block_math, font_size_pt=12
                    )
                return
        for child in element.children:
            if isinstance(child, NavigableString):
                s = str(child)
                if s:
                    self._add_formatted_text(paragraph, s)
                continue
            if not isinstance(child, Tag):
                continue
            if child.name == 'img':
                self._emit_img_tag(paragraph, child)
                continue
            if child.name == 'br':
                paragraph.add_run().add_break()
                continue
            if child.name == 'div':
                cl = child.get('class') or []
                if isinstance(cl, str):
                    cl = [cl]
                is_block = any(
                    x in ('math-block', 'math-display', 'formula-image-block') for x in cl
                )
                # 已在居中段内嵌套块级 div 时不再新开段落，减少同图分段、叠盖
                if is_block and paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    self._fill_inline(paragraph, child)
                else:
                    np = self.document.add_paragraph()
                    if is_block:
                        np.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    self._fill_inline(np, child)
                continue
            if child.name in ('span', 'strong', 'b', 'em', 'i', 'u', 'sub', 'sup'):
                self._fill_inline(paragraph, child)
                continue
            # 其它标签（如 p、font）：展开子节点
            self._fill_inline(paragraph, child)

    @staticmethod
    def _escape_html_body(text: str) -> str:
        """与 splitter 题干 HTML 一致：正文内 & < >。"""
        t = (text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        return t

    def _synthetic_math_inline_span(self, latex: str) -> Optional[Tag]:
        """为纯文本中的 $...$ 构造与 content_html 同结构的 math-inline，供 OLE/OMML 插入。"""
        raw = (latex or "").strip()
        if not raw:
            return None
        attr = html_module.escape(raw, quote=True)
        body = self._escape_html_body(raw)
        frag = f'<span class="math-inline" data-latex="{attr}">${body}$</span>'
        soup = BeautifulSoup(frag, "html.parser")
        return soup.select_one("span.math-inline")

    def _append_latex_inline_best_effort(self, paragraph, latex: str, font_size_pt: float) -> None:
        """与 _fill_inline 相同引擎顺序（EXPORT_FORMULA_PRIORITY）；失败则斜体 LaTeX。"""
        span = self._synthetic_math_inline_span(latex)
        if span is None:
            return
        if self._try_formula_engines_order(paragraph, span):
            return
        run = paragraph.add_run(f" ${latex.strip()} ")
        run.italic = True
        run.font.size = Pt(font_size_pt)

    def _add_formatted_text(self, paragraph, text: str, font_size_pt: float = 12):
        """普通文本 + `$...$` 行内公式；公式走 MathType/OMML，与题干同源。"""
        parts = (text or "").split("$")
        for i, part in enumerate(parts):
            if i % 2 == 0:
                if part:
                    run = paragraph.add_run(part)
                    run.font.size = Pt(font_size_pt)
            else:
                if not (part or "").strip():
                    continue
                self._append_latex_inline_best_effort(paragraph, part, font_size_pt)

    def _add_options(self, options: List[Any]):
        """添加选项"""
        # 每行两个选项
        for i in range(0, len(options), 2):
            para = self.document.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.5)
            
            # 第一个选项
            opt1 = options[i]
            run = para.add_run(f"{opt1.label}. ")
            run.font.bold = True
            self._add_option_body(para, opt1)
            
            # 第二个选项（如果有）
            if i + 1 < len(options):
                opt2 = options[i + 1]
                # 添加制表符分隔
                para.add_run('\t\t')
                run = para.add_run(f"{opt2.label}. ")
                run.font.bold = True
                self._add_option_body(para, opt2)

    def _fill_option_root(self, paragraph, root: Tag) -> None:
        """选项根：顶层多段与题干一致分段。"""
        first_top = True
        current_para = paragraph
        start_new_para_before_next_flow = False
        opt_indent = paragraph.paragraph_format.left_indent

        for child in root.children:
            if isinstance(child, NavigableString):
                s = str(child)
                if s.strip():
                    if start_new_para_before_next_flow:
                        current_para = self.document.add_paragraph()
                        current_para.paragraph_format.left_indent = opt_indent
                        start_new_para_before_next_flow = False
                    self._add_formatted_text(current_para, s)
                    first_top = False
                continue
            if not isinstance(child, Tag):
                continue
            if child.name == 'p':
                target = paragraph if first_top else self.document.add_paragraph()
                target.paragraph_format.left_indent = opt_indent
                self._fill_inline(target, child)
                current_para = target
                first_top = False
                start_new_para_before_next_flow = True
            elif child.name == 'div':
                cl = child.get('class') or []
                if isinstance(cl, str):
                    cl = [cl]
                is_center = any(
                    x in ('math-block', 'math-display', 'formula-image-block') for x in cl
                )
                if is_center:
                    target = self.document.add_paragraph()
                    target.paragraph_format.left_indent = opt_indent
                    target.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    self._fill_inline(target, child)
                    start_new_para_before_next_flow = True
                elif first_top:
                    target = paragraph
                    self._fill_inline(target, child)
                    current_para = target
                    start_new_para_before_next_flow = False
                else:
                    target = self.document.add_paragraph()
                    target.paragraph_format.left_indent = opt_indent
                    self._fill_inline(target, child)
                    current_para = target
                    start_new_para_before_next_flow = False
                first_top = False
            else:
                if start_new_para_before_next_flow:
                    current_para = self.document.add_paragraph()
                    current_para.paragraph_format.left_indent = opt_indent
                    start_new_para_before_next_flow = False
                self._fill_inline(current_para, child)
                first_top = False

    def _add_option_body(self, paragraph, opt: Any) -> None:
        oh = (getattr(opt, 'content_html', None) or '').strip()
        if oh:
            oh = html_module.unescape(oh)
            oh = oh.replace('\\"', '"').replace("\\'", "'").replace("\\/", "/")
            wrapped = BeautifulSoup(f'<div class="opt-root">{oh}</div>', 'html.parser')
            root = wrapped.select_one('.opt-root')
            if root:
                for span in root.select('span.option-label'):
                    span.decompose()
                self._fill_option_root(paragraph, root)
                if not re.search(r"<img\b", oh, re.I):
                    self._add_missing_images_inline(paragraph, getattr(opt, 'images', None) or [])
        else:
            self._add_formatted_text(paragraph, opt.content)
                
    def add_section_header(self, text: str):
        """添加章节标题（如"一、选择题"）"""
        para = self.document.add_paragraph()
        run = para.add_run(text)
        run.font.bold = True
        run.font.size = Pt(14)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)
        
    def save(self, output_path: str):
        """保存文档"""
        self._refresh_document_modified_time()
        self.document.save(output_path)

    def to_bytes(self) -> bytes:
        """导出为字节流"""
        self._refresh_document_modified_time()
        buffer = io.BytesIO()
        self.document.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()


def export_questions(questions: List[Question], output_path: str, 
                     title: str = "", watermark: str = None,
                     image_dir: Optional[str] = None,
                     include_answer: bool = False, 
                     include_analysis: bool = False) -> str:
    """
    便捷函数：导出题目到 Word
    
    Args:
        questions: 题目列表
        output_path: 输出路径
        title: 文档标题
        watermark: 水印文字
        include_answer: 是否包含答案
        include_analysis: 是否包含解析
        
    Returns:
        输出文件路径
    """
    exporter = WordExporter()
    exporter.create_document(title=title, watermark_text=watermark, image_dir=image_dir)

    nq = len(questions)
    seg_counts = [len(getattr(q, "content_export_segments", None) or []) for q in questions]
    n_with_seg = sum(1 for c in seg_counts if c > 0)
    logger.info(
        "Word 导出：共 %d 题，%d 题有 content_export_segments（为 0 的题干仍走 content_html，与按段修复无关）；"
        "各题 segment 条数=%s。若刚改 exporter 但行为不变，请重启 Parser（pnpm dev 默认 uvicorn 无热重载）。",
        nq,
        n_with_seg,
        seg_counts,
    )

    # 逐题导出：每题 deepcopy 题目对象，保证 segments/options 等嵌套结构与别题无共享引用；
    # add_question 内再清空 OMML 失败集、嵌图去重集、按 file_id 切换目录，避免跨题状态串扰。
    # （仍在同一 .docx 中追加段落；若需进程级隔离只能每题单文件再合并，成本高。）
    current_type = None
    for question in questions:
        q_local = copy.deepcopy(question)
        if q_local.type_name and q_local.type_name != current_type:
            exporter.add_section_header(f"{q_local.type_name}")
            current_type = q_local.type_name

        exporter.add_question(q_local, include_answer, include_analysis)
        
    exporter.save(output_path)
    return output_path
